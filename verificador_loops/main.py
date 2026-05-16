import os
import re
from collections import defaultdict
from docx import Document

# =========================================================
# CONFIG
# =========================================================

PASTA_ENTRADA = "entrada"
PASTA_SAIDA = "saida"

ARQUIVO_RELATORIO = os.path.join(
    PASTA_SAIDA,
    "relatorio.txt"
)

# =========================================================
# UTILS
# =========================================================

def normalizar(texto):

    texto = (
        texto.replace("\n", " ")
        .replace("\t", " ")
        .replace("–", "-")
        .strip()
        .upper()
    )

    # =====================================
    # REMOVE (VO)
    # =====================================

    texto = re.sub(
        r"\s*\(VO\)",
        "",
        texto
    )

    texto = texto.strip()

    # =====================================
    # NORMALIZA VOZERIO
    # =====================================

    if texto.startswith("VOZERIO"):
        return "VOZERIO"

    return texto


def eh_loop(texto):

    return bool(
        re.fullmatch(r"\d{3}", texto.strip())
    )


def eh_timecode(texto):

    return bool(
        re.fullmatch(r"\d{4}", texto.strip())
    )


def eh_personagem(texto):

    texto = normalizar(texto)

    if not texto:
        return False

    # LOOP
    if eh_loop(texto):
        return False

    # TIMECODE
    if eh_timecode(texto):
        return False

    # RUBRICA
    if texto.startswith("("):
        return False

    # CONTÉM TIMECODE
    if re.search(r"//\d{4}", texto):
        return False

    # MUITO GRANDE
    if len(texto) > 40:
        return False

    # PRECISA TER LETRAS
    if not re.search(r"[A-Z]", texto):
        return False

    return True


# =========================================================
# ARQUIVOS
# =========================================================

def encontrar_arquivos():

    marcado = None
    levantamento = None

    for arquivo in os.listdir(PASTA_ENTRADA):

        nome = arquivo.upper()

        if "_MARCADO" in nome and arquivo.endswith(".docx"):

            marcado = os.path.join(
                PASTA_ENTRADA,
                arquivo
            )

        if "_LEVANTAMENTO" in nome and arquivo.endswith(".docx"):

            levantamento = os.path.join(
                PASTA_ENTRADA,
                arquivo
            )

    return marcado, levantamento


# =========================================================
# EXTRAIR LINHAS DAS TABELAS
# =========================================================

def extrair_linhas_tabelas(doc_path):

    doc = Document(doc_path)

    linhas = []

    for tabela in doc.tables:

        for row in tabela.rows:

            linha = []

            for cell in row.cells:

                texto = normalizar(cell.text)

                linha.append(texto)

            linhas.append(linha)

    return linhas


# =========================================================
# MARCADO
# =========================================================

def analisar_marcado(doc_path):

    linhas = extrair_linhas_tabelas(doc_path)

    personagem_loops = defaultdict(set)

    loops_atuais = []

    regex_loop = re.compile(r"\d{3}")

    for linha in linhas:

        if not linha:
            continue

        # =====================================
        # COLUNA 0
        # =====================================

        col0 = normalizar(linha[0])

        if not col0:
            continue

        # =====================================
        # PROCURA TODOS OS LOOPS
        # =====================================

        loops_encontrados = regex_loop.findall(
            col0
        )

        # Se encontrou loops,
        # atualiza os loops atuais
        if loops_encontrados:

            loops_atuais = loops_encontrados
            continue

        # =====================================
        # PERSONAGEM
        # =====================================

        if eh_personagem(col0):

            for loop in loops_atuais:

                personagem_loops[
                    col0
                ].add(loop)

    # =====================================
    # RESULTADO
    # =====================================

    resultado = {}

    for personagem, loops in personagem_loops.items():

        resultado[personagem] = sorted(loops)

    return resultado


# =========================================================
# LEVANTAMENTO
# =========================================================

def analisar_levantamento(doc_path):

    doc = Document(doc_path)

    levantamento = {}

    regex_loop = re.compile(
        r"(?<!//)\b\d{3}\b"
    )
    regex_total = re.compile(r"//(\d+)")

    for tabela in doc.tables:

        for row in tabela.rows:

            if len(row.cells) < 2:
                continue

            # =====================================
            # PERSONAGEM
            # =====================================

            personagem = normalizar(
                row.cells[0].text
            )

            if not personagem:
                continue

            if not eh_personagem(personagem):
                continue

            # =====================================
            # TEXTO DOS LOOPS
            # =====================================

            loops_texto = normalizar(
                row.cells[1].text
            )

            # =====================================
            # EXTRAI LOOPS
            # =====================================

            loops = regex_loop.findall(
                loops_texto
            )

            # =====================================
            # EXTRAI TOTAL
            # =====================================

            total_match = regex_total.search(
                loops_texto
            )

            total_informado = None

            if total_match:

                total_informado = int(
                    total_match.group(1)
                )

            levantamento[personagem] = {
                "loops": sorted(set(loops)),
                "total_informado": total_informado
            }

    return levantamento


# =========================================================
# COMPARAÇÃO
# =========================================================

def comparar(marcado, levantamento):

    relatorio = []

    todos = sorted(
        set(marcado.keys()) |
        set(levantamento.keys())
    )

    for personagem in todos:

        loops_marcado = set(
            marcado.get(personagem, [])
        )

        levantamento_data = levantamento.get(
            personagem,
            {}
        )

        loops_levantamento = set(
            levantamento_data.get(
                "loops",
                []
            )
        )

        total_informado = levantamento_data.get(
            "total_informado"
        )

        # =====================================
        # FALTANDO
        # =====================================

        if personagem not in levantamento:

            relatorio.append(
                f"[FALTANDO] {personagem}\n"
                f"Somente no Marcado: "
                f"{' - '.join(sorted(loops_marcado))}"
            )

            continue

        # =====================================
        # EXTRA
        # =====================================

        if personagem not in marcado:

            relatorio.append(
                f"[EXTRA] {personagem}\n"
                f"Somente no Levantamento: "
                f"{' - '.join(sorted(loops_levantamento))}"
            )

            continue

        # =====================================
        # DIFERENÇAS
        # =====================================

        somente_marcado = sorted(
            loops_marcado - loops_levantamento
        )

        somente_levantamento = sorted(
            loops_levantamento - loops_marcado
        )

        erro = False

        texto_erro = (
            f"[ERRO] {personagem}\n"
        )

        # =====================================
        # LOOPS DIFERENTES
        # =====================================

        if (
            somente_marcado
            or
            somente_levantamento
        ):

            erro = True

            texto_erro += (
                f"Somente no Marcado: "
                f"{' - '.join(somente_marcado)}\n"
                f"Somente no Levantamento: "
                f"{' - '.join(somente_levantamento)}\n"
            )

        # =====================================
        # TOTAL INCORRETO
        # =====================================

        total_real = len(loops_levantamento)

        if (
            total_informado is not None
            and
            total_real != total_informado
        ):

            erro = True

            texto_erro += (
                f"Total informado: "
                f"{total_informado}\n"
                f"Total real: "
                f"{total_real}\n"
            )

        if erro:

            relatorio.append(
                texto_erro.strip()
            )

    return relatorio


# =========================================================
# SALVAR
# =========================================================

def salvar_relatorio(relatorio):

    os.makedirs(
        PASTA_SAIDA,
        exist_ok=True
    )

    with open(
        ARQUIVO_RELATORIO,
        "w",
        encoding="utf-8"
    ) as f:

        for linha in relatorio:

            f.write(linha + "\n\n")


# =========================================================
# MAIN
# =========================================================

def main():

    print("=" * 60)
    print("VERIFICADOR DE LEVANTAMENTO")
    print("=" * 60)

    marcado_path, levantamento_path = encontrar_arquivos()

    if not marcado_path:

        print()
        print("Arquivo _MARCADO.docx não encontrado.")
        return

    if not levantamento_path:

        print()
        print("Arquivo _LEVANTAMENTO.docx não encontrado.")
        return

    print()
    print("Arquivos encontrados:")
    print()

    print(
        f"Marcado: "
        f"{os.path.basename(marcado_path)}"
    )

    print(
        f"Levantamento: "
        f"{os.path.basename(levantamento_path)}"
    )

    print()
    print("Analisando marcado...")

    marcado = analisar_marcado(
        marcado_path
    )

    print(
        f"{len(marcado)} personagens encontrados."
    )

    print()
    print("Analisando levantamento...")

    levantamento = analisar_levantamento(
        levantamento_path
    )

    print(
        f"{len(levantamento)} personagens encontrados."
    )

    print()
    print("Comparando...")

    relatorio = comparar(
        marcado,
        levantamento
    )

    salvar_relatorio(relatorio)

    print()
    print("=" * 60)
    print("VERIFICAÇÃO CONCLUÍDA")
    print("=" * 60)

    print()
    print(
        f"Relatório salvo em:\n"
        f"{ARQUIVO_RELATORIO}"
    )


if __name__ == "__main__":
    main()